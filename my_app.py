from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
     pyttsx3.speak(text)

document = Document()

document.add_picture('pexels-photo-220453.jpeg', width = Inches(2.0))

name = input ('What is your name?')
speak('Hello' + name + 'how are you today')
Phone = input ('What is your number?')
Email = input ('What is your email?')

document.add_paragraph(  name + '|' + Email + '|' +  Phone)

document.add_heading ('About me')
document.add_paragraph(input ('Tell about yourself? '))

document.add_heading('Work Experiences')
p = document.add_paragraph()
company = input('Enter company name')
from_date = input('Enter date')
to_date = input('Enter date')

p.add_run(company + ' ').bold = True
p.add_run(from_date + ' - ' + to_date + '\n').italic = True

experience_details = input('Describe your experience at' + company)
p.add_run(experience_details)

while True:
    more_experience = input('Do you want to add more experience? say yes or no ')
    if more_experience.lower() == 'yes':
        p = document.add_paragraph()
        company = input('Enter company name')
        from_date = input('Enter date')
        to_date = input('Enter date')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + ' - ' + to_date + '\n').italic = True

        experience_details = input('Describe your experience at' + company)
        p.add_run(experience_details)
    else:
        break

document.add_heading( ' SKILLS').bold = True

skill_name = input('Enter skill name')
skill_experience_year = input('Number of years ')
b = document.add_paragraph(skill_name)
b.style = 'List Bullet'
while True:
    more_skills = input('Do you want to add more skills yes or no' )

    if more_skills.lower() == 'yes':
        skill_name = input('Enter skill name')
        skill_experience_year = input('Number of years ')
        b = document.add_paragraph(skill_name)
        b.style = 'List Bullet'
    else:
        break
    
document.save('cv.docx')